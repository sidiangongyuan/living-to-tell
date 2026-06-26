"""Articles feature: entry CRUD + tags + search.

Exposes the EntryRepository and SearchService from the existing writer backend.
"""
from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel, Field
from starlette.background import BackgroundTask

from deps import get_container
from writer.app.container import AppContainer
from writer.domain.enums import VersionType
from writer.domain.models.entry import Entry as DomainEntry
from writer.domain.models.entry_version import EntryVersion as DomainEntryVersion
from writer.services.epigraph import detect_epigraph, strip_epigraph

router = APIRouter(prefix="/api/articles", tags=["articles"])


# ---- DTOs matching frontend contract ----
class EntryOut(BaseModel):
    id: str
    title: str
    body: str
    entry_type: str
    created_at: Optional[str]
    updated_at: Optional[str]
    tags: list[str]
    archived_at: Optional[str]
    curation_status: str


class EntryCreate(BaseModel):
    title: str = ""
    body: str = ""
    tags: list[str] = Field(default_factory=list)


class EntryUpdate(BaseModel):
    title: str
    body: str
    tags: Optional[list[str]] = None


class EntryVersionOut(BaseModel):
    id: str
    entry_id: str
    version_type: str
    content: str
    title_snapshot: str
    tags: list[str]
    label: str
    reason: str
    word_count: int
    char_count: int
    created_at: Optional[str]
    provider: Optional[str]
    model: Optional[str]


class EntryVersionCreate(BaseModel):
    version_type: str = VersionType.MANUAL_CHECKPOINT.value
    label: str = ""
    provider: Optional[str] = None
    model: Optional[str] = None


class EntryVersionRestoreOut(BaseModel):
    entry: EntryOut
    snapshot_version_id: Optional[str]
    was_noop: bool


def _to_dto(e: DomainEntry) -> EntryOut:
    return EntryOut(
        id=e.id,
        title=e.title,
        body=e.body,
        entry_type=e.entry_type,
        created_at=e.created_at,
        updated_at=e.updated_at,
        tags=e.tags,
        archived_at=e.archived_at,
        curation_status=e.curation_status,
    )


def _parse_tags_snapshot(tags_snapshot: str) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for part in (tags_snapshot or "").split(","):
        tag = part.strip()
        if tag and tag.lower() not in seen:
            seen.add(tag.lower())
            result.append(tag)
    return result


def _version_to_dto(version: DomainEntryVersion) -> EntryVersionOut:
    return EntryVersionOut(
        id=version.id,
        entry_id=version.entry_id,
        version_type=version.version_type,
        content=version.content,
        title_snapshot=version.title_snapshot,
        tags=_parse_tags_snapshot(version.tags_snapshot),
        label=version.label,
        reason=version.reason,
        word_count=version.word_count,
        char_count=version.char_count,
        created_at=version.created_at,
        provider=version.provider,
        model=version.model,
    )


def _entry_or_404(entry_id: str, container: AppContainer) -> DomainEntry:
    entry = container.entry_repository.get(entry_id)
    if entry is None:
        raise HTTPException(404, "Entry not found")
    return entry


def _entry_export_filename(entry: DomainEntry, ext: str) -> str:
    title = (entry.title or "").strip() or "article"
    sanitized = "".join(ch for ch in title if ch not in '<>:"/\\|?*').strip()
    return f"{sanitized or 'article'}.{ext}"


def _export_entry_docx(entry: DomainEntry, output_path: str) -> str:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover - env-specific
        raise RuntimeError(
            "python-docx is required for DOCX export. Install it with: pip install python-docx"
        ) from exc

    document = Document()
    document.add_heading(entry.title.strip() or "Untitled", level=1)
    body = (entry.body or "").strip()
    if body:
        epigraph = detect_epigraph(body)
        rendered_body = strip_epigraph(body).rstrip() if epigraph is not None else body
        if epigraph is not None:
            for line in epigraph.quote.splitlines():
                para = document.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.35)
                para.paragraph_format.right_indent = Inches(0.55)
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run(line)
                run.italic = True

            attr_para = document.add_paragraph()
            attr_para.paragraph_format.right_indent = Inches(0.55)
            attr_para.paragraph_format.space_after = Pt(8)
            attr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            attr_para.add_run(epigraph.attribution)

        if rendered_body:
            for para in rendered_body.split("\n\n"):
                document.add_paragraph(para)
    document.save(output_path)
    return output_path


def _cleanup_temp_file(path: str) -> None:
    try:
        Path(path).unlink(missing_ok=True)
    except OSError:
        pass


# ---- Routes ----
@router.get("", response_model=list[EntryOut])
def list_entries(
    limit: int = Query(100, le=500),
    include_archived: bool = False,
    container: AppContainer = Depends(get_container),
) -> list[EntryOut]:
    entries = container.entry_repository.list_recent(
        limit=limit, include_archived=include_archived
    )
    return [_to_dto(e) for e in entries]


@router.get("/search", response_model=list[EntryOut])
def search_entries(
    q: str = Query(..., min_length=1),
    limit: int = Query(50, le=200),
    container: AppContainer = Depends(get_container),
) -> list[EntryOut]:
    results = container.search_service.search(q, limit=limit)
    return [_to_dto(e) for e in results]


@router.get("/{entry_id}", response_model=EntryOut)
def get_entry(
    entry_id: str,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    return _to_dto(_entry_or_404(entry_id, container))


@router.get("/{entry_id}/export")
def export_entry(
    entry_id: str,
    format: str = Query("md", pattern="^(txt|md|docx)$"),
    container: AppContainer = Depends(get_container),
):
    entry = _entry_or_404(entry_id, container)
    if format == "txt":
        content = container.text_exporter.export_entry(entry)
        return Response(content, media_type="text/plain; charset=utf-8")
    if format == "md":
        content = container.markdown_exporter.export_entry(entry)
        return Response(content, media_type="text/markdown; charset=utf-8")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    try:
        output = _export_entry_docx(entry, tmp.name)
    except Exception:
        _cleanup_temp_file(tmp.name)
        raise
    return FileResponse(
        output,
        filename=_entry_export_filename(entry, "docx"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(_cleanup_temp_file, output),
    )


@router.get("/{entry_id}/versions", response_model=list[EntryVersionOut])
def list_entry_versions(
    entry_id: str,
    container: AppContainer = Depends(get_container),
) -> list[EntryVersionOut]:
    _entry_or_404(entry_id, container)
    return [
        _version_to_dto(version)
        for version in container.version_history_service.list_history(entry_id)
    ]


@router.post("/{entry_id}/versions", response_model=EntryVersionOut, status_code=201)
def create_entry_version(
    entry_id: str,
    data: EntryVersionCreate,
    container: AppContainer = Depends(get_container),
) -> EntryVersionOut:
    _entry_or_404(entry_id, container)
    version_type = (data.version_type or VersionType.MANUAL_CHECKPOINT.value).strip()
    try:
        if version_type == VersionType.AI_BEFORE_APPLY.value:
            version = container.version_history_service.save_ai_before_apply(
                entry_id,
                label=data.label,
                provider=data.provider,
                model=data.model,
            )
        elif version_type == VersionType.MANUAL_CHECKPOINT.value:
            version = container.version_history_service.save_manual_checkpoint(
                entry_id,
                label=data.label,
            )
        else:
            raise ValueError("Unsupported version type")
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    return _version_to_dto(version)


@router.post("/{entry_id}/versions/{version_id}/restore", response_model=EntryVersionRestoreOut)
def restore_entry_version(
    entry_id: str,
    version_id: str,
    container: AppContainer = Depends(get_container),
) -> EntryVersionRestoreOut:
    try:
        outcome = container.version_history_service.restore(entry_id, version_id)
    except ValueError as exc:
        raise HTTPException(404, str(exc)) from exc
    entry = _entry_or_404(entry_id, container)
    return EntryVersionRestoreOut(
        entry=_to_dto(entry),
        snapshot_version_id=outcome.snapshot_version_id,
        was_noop=outcome.was_noop,
    )


@router.post("/{entry_id}/versions/{version_id}/clone", response_model=EntryOut, status_code=201)
def clone_entry_version(
    entry_id: str,
    version_id: str,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    _entry_or_404(entry_id, container)
    version = container.version_repository.get(version_id)
    if version is None or version.entry_id != entry_id:
        raise HTTPException(404, "Version not found")
    base_title = (version.title_snapshot or "").strip() or _entry_or_404(entry_id, container).title
    title = f"{base_title or '未命名文章'}（历史版本）"
    entry = container.entry_repository.create(
        title=title,
        body=version.content,
        tags=_parse_tags_snapshot(version.tags_snapshot),
    )
    return _to_dto(entry)


@router.delete("/{entry_id}/versions/{version_id}", status_code=204)
def delete_entry_version(
    entry_id: str,
    version_id: str,
    container: AppContainer = Depends(get_container),
):
    _entry_or_404(entry_id, container)
    try:
        container.version_history_service.delete_version(entry_id, version_id)
    except ValueError as exc:
        raise HTTPException(404, str(exc)) from exc


@router.post("", response_model=EntryOut, status_code=201)
def create_entry(
    data: EntryCreate,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    entry = container.entry_repository.create(
        title=data.title, body=data.body, tags=data.tags
    )
    return _to_dto(entry)


@router.put("/{entry_id}", response_model=EntryOut)
def update_entry(
    entry_id: str,
    data: EntryUpdate,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    entry = container.entry_repository.update(
        entry_id, title=data.title, body=data.body, tags=data.tags
    )
    if not entry:
        raise HTTPException(404, "Entry not found")
    return _to_dto(entry)


@router.delete("/{entry_id}", status_code=204)
def delete_entry(
    entry_id: str,
    container: AppContainer = Depends(get_container),
):
    success = container.entry_repository.delete(entry_id)
    if not success:
        raise HTTPException(404, "Entry not found")


@router.post("/{entry_id}/archive", response_model=EntryOut)
def archive_entry(
    entry_id: str,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    entry = container.entry_repository.archive(entry_id)
    if not entry:
        raise HTTPException(404, "Entry not found")
    return _to_dto(entry)


@router.post("/{entry_id}/unarchive", response_model=EntryOut)
def unarchive_entry(
    entry_id: str,
    container: AppContainer = Depends(get_container),
) -> EntryOut:
    entry = container.entry_repository.unarchive(entry_id)
    if not entry:
        raise HTTPException(404, "Entry not found")
    return _to_dto(entry)
