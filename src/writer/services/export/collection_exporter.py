"""Article collection exporters."""
from __future__ import annotations

from dataclasses import dataclass
from typing import Iterable, List, Optional

from writer.domain.models.collection import Collection
from writer.domain.models.collection_outline import CollectionOutlineItem
from writer.domain.models.entry import Entry
from writer.services.epigraph import detect_epigraph, strip_epigraph
from writer.storage.repositories.collection_outline_repository import (
    CollectionOutlineRepository,
)
from writer.storage.repositories.collection_repository import CollectionRepository
from writer.storage.repositories.entry_repository import EntryRepository


@dataclass
class _ArticleBundle:
    entry: Entry
    word_count: int


class CollectionExportService:
    def __init__(
        self,
        collection_repo: CollectionRepository,
        outline_repo: CollectionOutlineRepository | None = None,
        entry_repo: EntryRepository | None = None,
    ) -> None:
        self._collections = collection_repo
        self._outline = outline_repo
        self._entries = entry_repo

    def export_collection_txt(self, collection_id: str) -> str:
        collection, outline, entry_map, bundles = self._load_collection(collection_id)
        parts: List[str] = []
        title = collection.name or "Untitled collection"
        parts.append(f"{title}\n{'=' * max(1, len(title))}")
        if collection.description.strip():
            parts.append(collection.description.strip())
        if outline:
            parts.append(_render_outline_contents_text(outline, entry_map))
            parts.extend(_render_outline_text(outline, entry_map))
        elif bundles:
            toc = ["Contents", "--------"]
            for index, bundle in enumerate(bundles, start=1):
                article_title = bundle.entry.title or "Untitled article"
                toc.append(f"{index}. {article_title}  [{bundle.word_count} words]")
            parts.append("\n".join(toc))
            for bundle in bundles:
                parts.append(_render_article_text(bundle.entry))
        return "\n\n".join(parts).rstrip() + "\n"

    def export_collection_md(self, collection_id: str) -> str:
        collection, outline, entry_map, bundles = self._load_collection(collection_id)
        lines: List[str] = [f"# {collection.name or 'Untitled collection'}"]
        if collection.description.strip():
            lines.append("")
            lines.append(collection.description.strip())
        if outline:
            lines.extend(_render_outline_contents_markdown(outline, entry_map))
            lines.extend(_render_outline_markdown(outline, entry_map))
        elif bundles:
            lines.append("")
            lines.append("## Contents")
            for index, bundle in enumerate(bundles, start=1):
                article_title = bundle.entry.title or "Untitled article"
                lines.append(f"{index}. {article_title}  _({bundle.word_count} words)_")
            for bundle in bundles:
                lines.append("")
                lines.append(_render_article_markdown(bundle.entry, heading_level=2))
        return "\n".join(lines).rstrip() + "\n"

    def export_collection_docx(self, collection_id: str, output_path: str) -> str:
        collection, outline, entry_map, bundles = self._load_collection(collection_id)
        document = _new_document()
        document.add_heading(collection.name or "Untitled collection", level=0)
        if collection.description.strip():
            document.add_paragraph(collection.description.strip())
        if outline:
            document.add_heading("Contents", level=1)
            for line in _outline_content_lines(outline, entry_map):
                document.add_paragraph(line)
            _write_outline_docx(document, outline, entry_map)
        elif bundles:
            document.add_heading("Contents", level=1)
            for index, bundle in enumerate(bundles, start=1):
                title = bundle.entry.title or "Untitled article"
                document.add_paragraph(f"{index}. {title}  ({bundle.word_count} words)")
            for bundle in bundles:
                _write_article_docx(document, bundle.entry, heading_level=1)
        document.save(output_path)
        return output_path

    def _load_collection(
        self,
        collection_id: str,
    ) -> tuple[
        Collection,
        list[CollectionOutlineItem],
        dict[str, Entry],
        List[_ArticleBundle],
    ]:
        collection = self._collections.get(collection_id)
        if collection is None:
            raise ValueError(f"unknown collection: {collection_id!r}")
        bundles = [
            _ArticleBundle(entry=entry, word_count=_word_count(entry.body or ""))
            for entry in self._collections.list_entries(collection_id)
        ]
        outline = self._outline.list_for_collection(collection_id) if self._outline else []
        entry_ids = {item.entry_id for item in outline if item.entry_id}
        entry_map: dict[str, Entry] = {}
        if self._entries is not None:
            for entry_id in entry_ids:
                entry = self._entries.get(entry_id)
                if entry is not None:
                    entry_map[entry_id] = entry
        else:
            entry_map = {bundle.entry.id: bundle.entry for bundle in bundles if bundle.entry.id in entry_ids}

        # Backward compatibility: old collections with article order but no
        # linked outline should export exactly as before.
        linked_outline_count = sum(1 for item in outline if item.entry_id)
        effective_outline = outline if outline and (linked_outline_count or not bundles) else []
        return collection, effective_outline, entry_map, bundles


def _render_article_text(entry: Entry) -> str:
    title = entry.title or "Untitled article"
    body = (entry.body or "").strip()
    if not body:
        return f"{title}\n{'-' * max(1, len(title))}"
    return f"{title}\n{'-' * max(1, len(title))}\n\n{body}".rstrip()


def _outline_children(
    outline: Iterable[CollectionOutlineItem],
) -> dict[Optional[str], list[CollectionOutlineItem]]:
    children: dict[Optional[str], list[CollectionOutlineItem]] = {}
    for item in outline:
        children.setdefault(item.parent_id, []).append(item)
    for rows in children.values():
        rows.sort(key=lambda item: (item.sort_order, item.created_at or "", item.id))
    return children


def _walk_outline(
    outline: list[CollectionOutlineItem],
) -> list[tuple[CollectionOutlineItem, int]]:
    children = _outline_children(outline)
    result: list[tuple[CollectionOutlineItem, int]] = []
    seen: set[str] = set()

    def visit(parent_id: Optional[str], depth: int) -> None:
        for item in children.get(parent_id, []):
            if item.id in seen:
                continue
            seen.add(item.id)
            result.append((item, depth))
            visit(item.id, depth + 1)

    visit(None, 0)
    for item in sorted(outline, key=lambda row: (row.sort_order, row.created_at or "", row.id)):
        if item.id not in seen:
            seen.add(item.id)
            result.append((item, 0))
            visit(item.id, 1)
    return result


def _outline_content_lines(
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> list[str]:
    lines: list[str] = []
    for item, depth in _walk_outline(outline):
        if item.item_type == "note":
            continue
        title = item.title or entry_map.get(item.entry_id or "") or "Untitled"
        if not isinstance(title, str):
            title = title.title or "Untitled"
        words = ""
        if item.entry_id and item.entry_id in entry_map:
            words = f"  ({_word_count(entry_map[item.entry_id].body or '')} words)"
        lines.append(f"{'  ' * depth}- {title}{words}")
    return lines


def _render_outline_contents_text(
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> str:
    lines = ["Contents", "--------"]
    lines.extend(_outline_content_lines(outline, entry_map))
    return "\n".join(lines)


def _render_outline_contents_markdown(
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> list[str]:
    lines = ["", "## Contents"]
    lines.extend(_outline_content_lines(outline, entry_map))
    return lines


def _render_outline_text(
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> list[str]:
    rendered: list[str] = []
    for item, depth in _walk_outline(outline):
        if item.item_type == "note" and not item.entry_id:
            continue
        entry = entry_map.get(item.entry_id or "")
        title = item.title or entry.title if entry else item.title or "Untitled"
        underline = "=" if depth == 0 else "-"
        body = _render_entry_body_text(entry) if entry is not None else ""
        section = f"{title}\n{underline * max(1, len(title))}"
        if body:
            section = f"{section}\n\n{body}"
        rendered.append(section.rstrip())
    return rendered


def _render_outline_markdown(
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> list[str]:
    rendered: list[str] = []
    for item, depth in _walk_outline(outline):
        if item.item_type == "note" and not item.entry_id:
            continue
        entry = entry_map.get(item.entry_id or "")
        title = item.title or entry.title if entry else item.title or "Untitled"
        heading_level = max(2, min(5, depth + 2))
        rendered.append("")
        rendered.append(f"{'#' * heading_level} {title}")
        if entry is not None:
            body = _render_entry_body_markdown(entry)
            if body:
                rendered.append("")
                rendered.append(body)
    return rendered


def _write_outline_docx(
    document,
    outline: list[CollectionOutlineItem],
    entry_map: dict[str, Entry],
) -> None:
    for item, depth in _walk_outline(outline):
        if item.item_type == "note" and not item.entry_id:
            continue
        entry = entry_map.get(item.entry_id or "")
        title = item.title or entry.title if entry else item.title or "Untitled"
        heading_level = max(1, min(4, depth + 1))
        document.add_heading(title, level=heading_level)
        if entry is not None:
            _write_entry_body_docx(document, entry)


def _render_article_markdown(entry: Entry, *, heading_level: int) -> str:
    level = max(1, min(5, heading_level))
    title = entry.title or "Untitled article"
    body = (entry.body or "").strip()
    lines: List[str] = [f"{'#' * level} {title}"]
    if not body:
        return "\n".join(lines)
    epigraph = detect_epigraph(body)
    rendered_body = strip_epigraph(body).rstrip() if epigraph is not None else body
    lines.append("")
    if epigraph is not None:
        lines.extend(_render_epigraph_markdown_lines(epigraph))
        if rendered_body:
            lines.append("")
    if rendered_body:
        lines.append(rendered_body)
    return "\n".join(lines).rstrip()


def _render_entry_body_text(entry: Entry | None) -> str:
    if entry is None:
        return ""
    return (entry.body or "").strip()


def _render_entry_body_markdown(entry: Entry) -> str:
    body = (entry.body or "").strip()
    if not body:
        return ""
    epigraph = detect_epigraph(body)
    rendered_body = strip_epigraph(body).rstrip() if epigraph is not None else body
    lines: list[str] = []
    if epigraph is not None:
        lines.extend(_render_epigraph_markdown_lines(epigraph))
        if rendered_body:
            lines.append("")
    if rendered_body:
        lines.append(rendered_body)
    return "\n".join(lines).rstrip()


def _write_article_docx(document, entry: Entry, *, heading_level: int) -> None:
    document.add_heading(entry.title or "Untitled article", level=heading_level)
    _write_entry_body_docx(document, entry)


def _write_entry_body_docx(document, entry: Entry) -> None:
    body = (entry.body or "").strip()
    if not body:
        return
    epigraph = detect_epigraph(body)
    rendered_body = strip_epigraph(body).rstrip() if epigraph is not None else body
    if epigraph is not None:
        _write_epigraph_docx(document, epigraph)
    if rendered_body:
        for para in rendered_body.split("\n\n"):
            document.add_paragraph(para)


def _word_count(body: str) -> int:
    if not body:
        return 0
    total = 0
    buf: list[str] = []
    for ch in body:
        if "\u4e00" <= ch <= "\u9fff" or "\u3400" <= ch <= "\u4dbf":
            if buf:
                if "".join(buf).strip():
                    total += 1
                buf = []
            total += 1
        else:
            buf.append(ch)
    if buf:
        total += len([token for token in "".join(buf).split() if token.strip()])
    return total


def _new_document():
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - env-specific
        raise RuntimeError(
            "python-docx is required for DOCX export. Install it with: pip install python-docx"
        ) from exc
    return Document()


def _render_epigraph_markdown_lines(epigraph) -> list[str]:
    lines = [f"> {line}" if line else ">" for line in epigraph.quote.splitlines()]
    lines.append(">")
    lines.append(f"> -- {epigraph.attribution}")
    return lines


def _write_epigraph_docx(document, epigraph) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore

    for line in epigraph.quote.splitlines():
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.35)
        para.paragraph_format.right_indent = Inches(0.55)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.italic = True

    attr_para = document.add_paragraph()
    attr_para.paragraph_format.right_indent = Inches(0.55)
    attr_para.paragraph_format.space_after = Pt(8)
    attr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    attr_para.add_run(epigraph.attribution)
