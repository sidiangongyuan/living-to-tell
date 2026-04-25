"""Work & Collection exporters (M8).

Three formats: plain text, Markdown, and DOCX. The DOCX writer uses
``python-docx`` (added in M8). All exporters render only sections in
their current ``sort_order`` and never include version-history snapshots.

A section's ``section_type``:
* ``heading``  → emitted as a level-2 heading (``##`` / ``Heading 2``).
* ``body``     → emitted as a paragraph.

Collection exports include a small auto-generated table of contents
(work titles + statuses + word counts) followed by each work in
collection order.
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import List, Optional

from writer.domain.enums import SectionType
from writer.domain.models.collection import Collection
from writer.domain.models.work import Work
from writer.domain.models.work_section import WorkSection
from writer.services.work_service import _word_count
from writer.storage.repositories.collection_repository import (
    CollectionRepository,
)
from writer.storage.repositories.work_repository import WorkRepository
from writer.storage.repositories.work_section_repository import (
    WorkSectionRepository,
)


@dataclass
class _WorkBundle:
    work: Work
    sections: List[WorkSection]
    word_count: int


class WorkExportService:
    def __init__(
        self,
        work_repo: WorkRepository,
        section_repo: WorkSectionRepository,
        collection_repo: CollectionRepository,
    ) -> None:
        self._works = work_repo
        self._sections = section_repo
        self._collections = collection_repo

    # ------------------------------------------------------------------
    # Single work
    # ------------------------------------------------------------------
    def export_work_txt(self, work_id: str) -> str:
        bundle = self._load_work(work_id)
        return self._render_work_text(bundle) + "\n"

    def export_work_md(self, work_id: str) -> str:
        bundle = self._load_work(work_id)
        return self._render_work_markdown(bundle, base_heading=1) + "\n"

    def export_work_docx(self, work_id: str, output_path: str) -> str:
        bundle = self._load_work(work_id)
        document = _new_document()
        self._write_work_docx(document, bundle, base_heading=1)
        document.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Collection
    # ------------------------------------------------------------------
    def export_collection_txt(self, collection_id: str) -> str:
        collection, bundles = self._load_collection(collection_id)
        parts: List[str] = []
        header = collection.name or "Untitled collection"
        parts.append(f"{header}\n{'=' * max(1, len(header))}")
        if collection.description.strip():
            parts.append(collection.description.strip())

        toc_lines = ["Contents", "--------"]
        for idx, bundle in enumerate(bundles, start=1):
            toc_lines.append(
                f"{idx}. {bundle.work.title or 'Untitled work'}  "
                f"[{bundle.work.status}, {bundle.word_count} words]"
            )
        parts.append("\n".join(toc_lines))

        for bundle in bundles:
            parts.append(self._render_work_text(bundle))
        return "\n\n".join(parts).rstrip() + "\n"

    def export_collection_md(self, collection_id: str) -> str:
        collection, bundles = self._load_collection(collection_id)
        lines: List[str] = [f"# {collection.name or 'Untitled collection'}"]
        if collection.description.strip():
            lines.append("")
            lines.append(collection.description.strip())
        lines.append("")
        lines.append("## Contents")
        for idx, bundle in enumerate(bundles, start=1):
            title = bundle.work.title or "Untitled work"
            lines.append(
                f"{idx}. {title}  _({bundle.work.status}, "
                f"{bundle.word_count} words)_"
            )
        for bundle in bundles:
            lines.append("")
            lines.append(self._render_work_markdown(bundle, base_heading=2))
        return "\n".join(lines).rstrip() + "\n"

    def export_collection_docx(
        self, collection_id: str, output_path: str
    ) -> str:
        collection, bundles = self._load_collection(collection_id)
        document = _new_document()
        document.add_heading(
            collection.name or "Untitled collection", level=0
        )
        if collection.description.strip():
            document.add_paragraph(collection.description.strip())
        document.add_heading("Contents", level=1)
        for idx, bundle in enumerate(bundles, start=1):
            title = bundle.work.title or "Untitled work"
            document.add_paragraph(
                f"{idx}. {title}  ({bundle.work.status}, "
                f"{bundle.word_count} words)"
            )
        for bundle in bundles:
            self._write_work_docx(document, bundle, base_heading=1)
        document.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Renderers
    # ------------------------------------------------------------------
    @staticmethod
    def _render_work_text(bundle: _WorkBundle) -> str:
        title = bundle.work.title or "Untitled work"
        parts: List[str] = [f"{title}\n{'=' * max(1, len(title))}"]
        if bundle.work.summary.strip():
            parts.append(bundle.work.summary.strip())
        for section in bundle.sections:
            content = (section.content or "").strip()
            if section.section_type == SectionType.HEADING.value:
                heading = content or "Untitled section"
                parts.append(
                    f"{heading}\n{'-' * max(1, len(heading))}"
                )
            elif content:
                parts.append(content)
        return "\n\n".join(parts).rstrip()

    @staticmethod
    def _render_work_markdown(
        bundle: _WorkBundle, *, base_heading: int
    ) -> str:
        base = max(1, min(5, base_heading))
        title = bundle.work.title or "Untitled work"
        lines: List[str] = [f"{'#' * base} {title}"]
        if bundle.work.summary.strip():
            lines.append("")
            lines.append(bundle.work.summary.strip())
        for section in bundle.sections:
            content = (section.content or "").strip()
            if section.section_type == SectionType.HEADING.value:
                heading = content or "Untitled section"
                lines.append("")
                lines.append(f"{'#' * (base + 1)} {heading}")
            elif content:
                lines.append("")
                lines.append(content)
        return "\n".join(lines).rstrip()

    @staticmethod
    def _write_work_docx(
        document, bundle: _WorkBundle, *, base_heading: int
    ) -> None:
        title = bundle.work.title or "Untitled work"
        document.add_heading(title, level=base_heading)
        if bundle.work.summary.strip():
            document.add_paragraph(bundle.work.summary.strip())
        for section in bundle.sections:
            content = (section.content or "").strip()
            if section.section_type == SectionType.HEADING.value:
                document.add_heading(
                    content or "Untitled section",
                    level=min(9, base_heading + 1),
                )
            elif content:
                # Preserve blank lines between paragraphs.
                for para in content.split("\n\n"):
                    document.add_paragraph(para)

    # ------------------------------------------------------------------
    # Loaders
    # ------------------------------------------------------------------
    def _load_work(self, work_id: str) -> _WorkBundle:
        work = self._works.get(work_id)
        if work is None:
            raise ValueError(f"unknown work: {work_id!r}")
        sections = self._sections.list_for_work(work_id)
        wc = sum(_word_count(s.content) for s in sections)
        return _WorkBundle(work=work, sections=sections, word_count=wc)

    def _load_collection(
        self, collection_id: str
    ) -> tuple[Collection, List[_WorkBundle]]:
        collection = self._collections.get(collection_id)
        if collection is None:
            raise ValueError(f"unknown collection: {collection_id!r}")
        bundles: List[_WorkBundle] = []
        for work in self._collections.list_works(collection_id):
            sections = self._sections.list_for_work(work.id)
            wc = sum(_word_count(s.content) for s in sections)
            bundles.append(
                _WorkBundle(work=work, sections=sections, word_count=wc)
            )
        return collection, bundles


# ---------------------------------------------------------------------------
# python-docx is imported lazily so the rest of the app can still load on
# environments where the wheel is not available (e.g. minimal CI image).
# ---------------------------------------------------------------------------
def _new_document():
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - env-specific
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        ) from exc
    return Document()
